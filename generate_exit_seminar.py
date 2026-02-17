#!/usr/bin/env python3
"""Generate Exit Seminar Slideshow Content as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import sys

print("[1/12] Creating document structure...", flush=True)
doc = Document()

# Styles setup
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper functions
def add_slide_header(doc, slide_num, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f'SLIDE {slide_num}: {title}')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)
    if subtitle:
        p2 = doc.add_paragraph()
        run2 = p2.add_run(subtitle)
        run2.italic = True
        run2.font.size = Pt(12)
        run2.font.color.rgb = RGBColor(100, 100, 100)

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f'[Speaker Notes: {text}]')
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 128, 0)

def add_divider(doc):
    p = doc.add_paragraph()
    run = p.add_run('_' * 80)
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.font.size = Pt(8)

# ============================================================
# TITLE PAGE
# ============================================================
print("[2/12] Writing title slide...", flush=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('EXIT SEMINAR SLIDESHOW CONTENT')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PhD Exit Seminar Presentation Guide')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nDynamic Pathfinding for Autonomous Systems:\nAn Efficient Grid-Map Framework for Classical Search')
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nElshahed Amr Moustafa Mohamed Aly Elsayed\nDoctor of Philosophy\nUniversiti Sains Malaysia\n2026')
run.font.size = Pt(12)

doc.add_paragraph('\nThis document contains all content needed for each slide of your exit seminar presentation. '
                   'Copy the content into your preferred presentation software (PowerPoint, Google Slides, LaTeX Beamer).')

doc.add_page_break()

# ============================================================
# SLIDE 1: TITLE SLIDE
# ============================================================
add_slide_header(doc, 1, 'TITLE SLIDE')
add_bullet(doc, 'Title: Dynamic Pathfinding for Autonomous Systems: An Efficient Grid-Map Framework for Classical Search')
add_bullet(doc, 'Candidate: Elshahed Amr Moustafa Mohamed Aly Elsayed')
add_bullet(doc, 'Degree: Doctor of Philosophy')
add_bullet(doc, 'Institution: Universiti Sains Malaysia')
add_bullet(doc, 'Year: 2026')
add_bullet(doc, 'Supervisor(s): [Add supervisor name(s)]')
add_note(doc, 'Welcome the audience. State your name, thesis title, and thank the committee for attending.')
add_divider(doc)

# ============================================================
# SLIDE 2: PRESENTATION OUTLINE
# ============================================================
print("[3/12] Writing outline slides...", flush=True)
add_slide_header(doc, 2, 'PRESENTATION OUTLINE')
add_bullet(doc, '1. Introduction & Motivation')
add_bullet(doc, '2. Problem Statement')
add_bullet(doc, '3. Research Questions & Objectives')
add_bullet(doc, '4. Literature Review Summary')
add_bullet(doc, '5. Methodology: Research Framework')
add_bullet(doc, '6. Methodology: ILS Framework')
add_bullet(doc, '7. Methodology: AILS Framework')
add_bullet(doc, '8. Results: ILS Experiments')
add_bullet(doc, '9. Results: AILS Experiments')
add_bullet(doc, '10. Results: Ablation Study')
add_bullet(doc, '11. Cross-Study Discussion')
add_bullet(doc, '12. Contributions & Significance')
add_bullet(doc, '13. Limitations & Future Work')
add_bullet(doc, '14. Conclusion')
add_bullet(doc, '15. Q&A')
add_note(doc, 'Briefly walk through the outline to set expectations. Estimated total time: 30-45 minutes for presentation, 15-30 minutes for Q&A.')
add_divider(doc)

# ============================================================
# SLIDE 3: INTRODUCTION & MOTIVATION
# ============================================================
add_slide_header(doc, 3, 'INTRODUCTION & MOTIVATION', 'Why does efficient pathfinding matter for biosecurity?')

p = doc.add_paragraph()
run = p.add_run('Background Context:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Autonomous systems (drones, robots, vehicles) are increasingly deployed in biosecurity-sensitive environments')
add_bullet(doc, 'Key environments: agricultural facilities, healthcare settings, transportation hubs, ports, border control')
add_bullet(doc, 'These systems need to compute collision-free routes quickly while accounting for spatially varying risk')
add_bullet(doc, 'The COVID-19 pandemic demonstrated the devastating impact of biological threats (millions of lives, trillions in economic damage)')

p = doc.add_paragraph()
run = p.add_run('The Pathfinding Challenge:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Grid-based pathfinding is the standard approach for autonomous navigation')
add_bullet(doc, 'Standard algorithms (A*, Dijkstra) explore the ENTIRE grid -- slow on large or cluttered maps')
add_bullet(doc, 'Risk-annotated grids add complexity: each cell has both traversability AND exposure risk')
add_bullet(doc, 'Real-time constraints: contamination surveillance, emergency response, rapid deployment')

p = doc.add_paragraph()
run = p.add_run('Key Insight:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'In most navigation scenarios, the optimal path stays close to the straight line between start and goal')
add_bullet(doc, 'What if we could restrict the search to just a narrow corridor around that straight line?')

add_note(doc, 'Use a visual example: show a grid map with start/goal, and illustrate how A* explores the entire grid while ILS focuses on a narrow band. This sets up the core idea of the thesis.')
add_divider(doc)

# ============================================================
# SLIDE 4: BIOSECURITY CONTEXT
# ============================================================
add_slide_header(doc, 4, 'BIOSECURITY CONTEXT', 'Malaysia and the global biosecurity landscape')

p = doc.add_paragraph()
run = p.add_run('Global Biosecurity Framework:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'WHO International Health Regulations: binding obligations for disease detection and response')
add_bullet(doc, 'Biosecurity spans: agricultural, public health, environmental, laboratory, and cyber domains')
add_bullet(doc, 'Cybersecurity underpins all modern biosecurity systems (surveillance, data, communication)')

p = doc.add_paragraph()
run = p.add_run('Malaysian Context:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Strategic position along the Strait of Malacca (~25% of global maritime trade)')
add_bullet(doc, 'Cybersecurity Act 2024: comprehensive framework for critical infrastructure')
add_bullet(doc, 'CAAM drone regulations: categorized by weight, commercial permits required')
add_bullet(doc, 'Exposure to biosecurity risks from endemic diseases, imported pathogens, potential bioterrorism')

p = doc.add_paragraph()
run = p.add_run('[Visual: Include the biosecurity_types_framework.png figure showing biosecurity domains]')
run.italic = True
run.font.color.rgb = RGBColor(0, 0, 200)

add_note(doc, 'Keep this slide brief (2 minutes max). The audience wants to hear about the algorithms, not a lecture on biosecurity policy. Just establish that the application context is real and important.')
add_divider(doc)

# ============================================================
# SLIDE 5: PROBLEM STATEMENT
# ============================================================
print("[4/12] Writing problem statement slides...", flush=True)
add_slide_header(doc, 5, 'PROBLEM STATEMENT')

p = doc.add_paragraph()
run = p.add_run('Problem 1: Real-time Computational Constraints')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(153, 0, 0)

add_bullet(doc, 'State space grows QUADRATICALLY with grid dimensions')
add_bullet(doc, 'Risk-annotated grids require evaluation of composite cost functions at every cell')
add_bullet(doc, 'Cluttered environments with narrow aisles force exploration of many dead-end routes')
add_bullet(doc, 'Computational delays undermine mission effectiveness in time-critical biosecurity operations')

p = doc.add_paragraph()
run = p.add_run('Problem 2: Limited Adaptability to Dynamic Environments')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(153, 0, 0)

add_bullet(doc, 'Real-world environments change: moving obstacles, updated risk maps, new quarantine zones')
add_bullet(doc, 'Most systems restart planning from scratch -- discarding all previous computation')
add_bullet(doc, 'Complete re-planning is a critical bottleneck for frequent updates')
add_bullet(doc, 'No mechanism to efficiently update plans with only locally relevant changes')

add_note(doc, 'Emphasize the GAP: existing methods either work on the full grid (slow) or require specific cost models (JPS for uniform costs only). There is no corridor-based method for weighted/risk grids.')
add_divider(doc)

# ============================================================
# SLIDE 6: RESEARCH QUESTIONS & OBJECTIVES
# ============================================================
add_slide_header(doc, 6, 'RESEARCH QUESTIONS & OBJECTIVES')

p = doc.add_paragraph()
run = p.add_run('Research Question 1 (RQ1):')
run.bold = True
run.font.size = Pt(12)
doc.add_paragraph('To what extent can constraining search to an ILS corridor reduce computational requirements while preserving path optimality?')

p = doc.add_paragraph()
run = p.add_run('Research Hypothesis 1 (RH1):')
run.bold = True
run.font.size = Pt(12)
add_bullet(doc, 'Expected 40-70% reduction in runtime and node expansions')
add_bullet(doc, 'Path lengths within 5% of optimal solutions')

p = doc.add_paragraph()
run = p.add_run('Research Question 2 (RQ2):')
run.bold = True
run.font.size = Pt(12)
doc.add_paragraph('Can an adaptive corridor mechanism that dynamically adjusts width based on local obstructions support efficient planning across diverse environments?')

p = doc.add_paragraph()
run = p.add_run('Research Hypothesis 2 (RH2):')
run.bold = True
run.font.size = Pt(12)
add_bullet(doc, 'Maintain 50-80% of ILS speedup under moderate dynamics')
add_bullet(doc, 'Re-planning latencies 3-5x faster than complete re-planning')

p = doc.add_paragraph()
run = p.add_run('Objectives:')
run.bold = True
run.font.size = Pt(13)
add_bullet(doc, 'O1: ', bold_prefix='')
doc.add_paragraph('O1: Design and evaluate the Incremental Line Search (ILS) framework')
doc.add_paragraph('O2: Develop and validate the Adaptive ILS (AILS) corridor control mechanism')

add_note(doc, 'Highlight that each RQ maps to one objective and one contribution. The structure is clean and traceable.')
add_divider(doc)

# ============================================================
# SLIDE 7: LITERATURE REVIEW SUMMARY
# ============================================================
print("[5/12] Writing literature review slides...", flush=True)
add_slide_header(doc, 7, 'LITERATURE REVIEW SUMMARY', 'Key areas and identified gaps')

p = doc.add_paragraph()
run = p.add_run('Key Areas Reviewed:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Classical search: A*, Dijkstra, BFS -- well-understood but explore full grid')
add_bullet(doc, 'Heuristic & weighted search: Weighted A*, Theta*, ANYA* -- speed/quality tradeoffs')
add_bullet(doc, 'Incremental replanning: D* Lite, LPA*, ARA* -- efficient for dynamic updates but full-grid memory')
add_bullet(doc, 'Symmetry reduction: JPS -- dramatic speedup but ONLY uniform-cost grids')
add_bullet(doc, 'Learning-based: Neural heuristics, RL policies -- lack formal guarantees')
add_bullet(doc, 'Subgoal graphs: Preprocessing-based -- fast queries but static environments only')

p = doc.add_paragraph()
run = p.add_run('Research Gap 1:')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(153, 0, 0)
doc.add_paragraph('No corridor-constrained search method exists for heterogeneous risk-weighted grids. JPS requires uniform costs. Subgoal methods require static preprocessing.')

p = doc.add_paragraph()
run = p.add_run('Research Gap 2:')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(153, 0, 0)
doc.add_paragraph('No adaptive mechanism dynamically adjusts search scope based on local obstructions. D* Lite maintains full-grid structures. Fixed-bound methods do not adapt locally.')

p = doc.add_paragraph()
run = p.add_run('[Visual: Table showing Gap-Objective alignment: G1 -> O1 -> ILS, G2 -> O2 -> AILS]')
run.italic = True
run.font.color.rgb = RGBColor(0, 0, 200)

add_note(doc, 'Keep this concise. The literature review supports two clear gaps that map directly to your two contributions. Do not try to cover every paper.')
add_divider(doc)

# ============================================================
# SLIDE 8-9: METHODOLOGY OVERVIEW
# ============================================================
print("[6/12] Writing methodology slides...", flush=True)
add_slide_header(doc, 8, 'METHODOLOGY: RESEARCH FRAMEWORK', 'Three-pipeline progressive design')

p = doc.add_paragraph()
run = p.add_run('Research Framework Overview:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Three pipelines sharing common grid representation, preprocessing, and evaluation metrics:')
add_bullet(doc, 'Pipeline 1 (Baseline): Full-grid search with 5 classical algorithms (A*, Dijkstra, BFS, DFS, Best-First)', level=1)
add_bullet(doc, 'Pipeline 2 (ILS-Enhanced): Uniform-width corridor + incremental expansion', level=1)
add_bullet(doc, 'Pipeline 3 (AILS): Density-adaptive corridor via integral images', level=1)

p = doc.add_paragraph()
run = p.add_run('Four Datasets:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'DS1: 6,000 synthetic 200x200 grids at 10%, 20%, 30% density')
add_bullet(doc, 'DS2: Variable-size grids (50x50 to 500x500), 10%-40% density')
add_bullet(doc, 'DS3: Five obstacle topologies (Random, Clustered, Maze, Room, Open)')
add_bullet(doc, 'DS4: Satellite-derived real-world grid')

p = doc.add_paragraph()
run = p.add_run('[Visual: Use the framework_overview.png figure from the thesis]')
run.italic = True
run.font.color.rgb = RGBColor(0, 0, 200)

add_note(doc, 'Explain the progressive design: Baseline establishes reference, ILS adds corridor, AILS adds adaptivity. Each pipeline adds one layer of sophistication.')
add_divider(doc)

# ============================================================
# SLIDE 9: ILS FRAMEWORK
# ============================================================
add_slide_header(doc, 9, 'METHODOLOGY: ILS FRAMEWORK', 'Incremental Line Search -- the core idea')

p = doc.add_paragraph()
run = p.add_run('How ILS Works (step by step):')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Step 1: Compute Bresenham line from start to goal (integer arithmetic, O(L) time)')
add_bullet(doc, 'Step 2: Build a uniform-width corridor of width w0 around the Bresenham line')
add_bullet(doc, 'Step 3: Run ANY classical search algorithm inside the corridor only')
add_bullet(doc, 'Step 4: If path found --> return it (with line-of-sight post-processing for A*/Dijkstra)')
add_bullet(doc, 'Step 5: If no path found --> widen corridor by delta_w, go to Step 3')
add_bullet(doc, 'Step 6: If corridor reaches full grid width --> report failure')

p = doc.add_paragraph()
run = p.add_run('Key Properties:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Algorithm-agnostic: Works as a WRAPPER around any search algorithm')
add_bullet(doc, 'Any cost model: Uniform, weighted, risk-annotated grids all supported')
add_bullet(doc, 'Complete: Fallback expansion guarantees path is found if one exists')
add_bullet(doc, 'Within-corridor optimal: Optimal algorithms remain optimal within the corridor')
add_bullet(doc, 'No preprocessing: No offline computation required')

p = doc.add_paragraph()
run = p.add_run('[Visual: Animation/diagram showing corridor construction around Bresenham line, then search within corridor]')
run.italic = True
run.font.color.rgb = RGBColor(0, 0, 200)

add_note(doc, 'This is the CORE slide. Spend 3-4 minutes here. Use a visual to show the Bresenham line, the corridor being built, and the search happening inside it. Contrast with full-grid A* exploring the entire map.')
add_divider(doc)

# ============================================================
# SLIDE 10: AILS FRAMEWORK
# ============================================================
add_slide_header(doc, 10, 'METHODOLOGY: AILS FRAMEWORK', 'Adaptive Incremental Line Search')

p = doc.add_paragraph()
run = p.add_run('Why Adaptive?')
run.bold = True
run.font.size = Pt(12)
doc.add_paragraph('ILS uses a UNIFORM-width corridor. Problem: in heterogeneous environments, the corridor must be wide enough for the densest segment, wasting space in open areas. AILS solves this with PER-POINT adaptive width.')

p = doc.add_paragraph()
run = p.add_run('Four-Stage Architecture:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Stage 1: Bresenham reference line generation')
add_bullet(doc, 'Stage 2: Per-point density estimation via integral image (O(1) per query)')
add_bullet(doc, 'Stage 3: Adaptive corridor construction using density-dependent radius formula:')
doc.add_paragraph('    r(p) = r_min + floor((r_max - r_min) * sigma(p)^alpha)')
add_bullet(doc, 'Stage 4: Corridor-constrained A* search with BFS fallback expansion')

p = doc.add_paragraph()
run = p.add_run('Three Strategies (auto-selected):')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Base: Fixed width when Bresenham line is obstacle-free (cheapest)')
add_bullet(doc, 'Standard: Density-adaptive when obstacles present, gradient small')
add_bullet(doc, 'Predictive: Gradient-enhanced when density changes rapidly (widens BEFORE dense regions)')

p = doc.add_paragraph()
run = p.add_run('[Visual: Use the ails_architecture.png and strategy_flowchart.png figures]')
run.italic = True
run.font.color.rgb = RGBColor(0, 0, 200)

add_note(doc, 'Explain with a visual: show how the corridor is narrow in open areas and wide near obstacles. Contrast uniform ILS corridor vs adaptive AILS corridor on the same map.')
add_divider(doc)

# ============================================================
# SLIDE 11: INTEGRAL IMAGE
# ============================================================
add_slide_header(doc, 11, 'KEY TECHNIQUE: INTEGRAL IMAGE', 'How O(1) density queries work')

p = doc.add_paragraph()
run = p.add_run('What is an Integral Image?')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'A 2D cumulative sum table (also called summed-area table)')
add_bullet(doc, 'I(x,y) = sum of all obstacle values in rectangle from (0,0) to (x,y)')
add_bullet(doc, 'Built in O(|V|) time -- single pass over the grid')

p = doc.add_paragraph()
run = p.add_run('How Density Queries Work:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'To count obstacles in ANY rectangular window:')
add_bullet(doc, 'count = I(x2,y2) - I(x1-1,y2) - I(x2,y1-1) + I(x1-1,y1-1)', level=1)
add_bullet(doc, 'Just 4 lookups + 3 arithmetic operations = O(1)')
add_bullet(doc, 'Works for any window size -- no iteration needed')

p = doc.add_paragraph()
run = p.add_run('Why This Matters for AILS:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'AILS queries density for every cell on the Bresenham line')
add_bullet(doc, 'Without integral image: O(omega^2) per query --> O(L * omega^2) total')
add_bullet(doc, 'With integral image: O(1) per query --> O(L) total')
add_bullet(doc, 'This makes density estimation essentially free compared to the search itself')

p = doc.add_paragraph()
run = p.add_run('[Visual: Diagram showing integral image construction and the 4-corner lookup formula]')
run.italic = True
run.font.color.rgb = RGBColor(0, 0, 200)

add_note(doc, 'This is a nice technical detail to explain clearly. The integral image is a well-known computer vision technique applied in a novel context here.')
add_divider(doc)

# ============================================================
# SLIDE 12-13: ILS RESULTS
# ============================================================
print("[7/12] Writing results slides...", flush=True)
add_slide_header(doc, 12, 'RESULTS: ILS EXPERIMENTS', 'DS1: 6,000 synthetic 200x200 grids')

p = doc.add_paragraph()
run = p.add_run('Headline Results:')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 102, 0)

add_bullet(doc, 'Average execution time reduction: 87.31% across all algorithms and densities')
add_bullet(doc, 'Average node reduction: 71.44%')
add_bullet(doc, 'All improvements statistically significant (p < 0.05)')

p = doc.add_paragraph()
run = p.add_run('By Algorithm (at 10% density):')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Best-First Search: 95.52% time reduction (highest single value)')
add_bullet(doc, 'A*: 94.81% time reduction')
add_bullet(doc, 'DFS: 92.33% time reduction')
add_bullet(doc, 'Dijkstra & BFS: >60% time reduction')

p = doc.add_paragraph()
run = p.add_run('Path Quality:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Optimal algorithms (A*, Dijkstra, BFS): path optimality PRESERVED')
add_bullet(doc, 'A* with line-of-sight: 69.54-86.37% Euclidean path improvement')
add_bullet(doc, 'DFS path length improved by up to 93.74% (corridor acts as guide)')
add_bullet(doc, 'Best-First: 63.24% average path improvement')

p = doc.add_paragraph()
run = p.add_run('Density Effect:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Improvements decrease with density (expected: more obstacles = more expansions)')
add_bullet(doc, 'A*: 94.81% at 10% density --> 82.77% at 30% density (still excellent)')

p = doc.add_paragraph()
run = p.add_run('Satellite Data (DS4): ILS maintained effectiveness on real-world grid with non-uniform obstacle distribution')
run.italic = True

p = doc.add_paragraph()
run = p.add_run('[Visual: Bar chart comparing standard vs ILS for each algorithm. Table of results.]')
run.italic = True
run.font.color.rgb = RGBColor(0, 0, 200)

add_note(doc, 'These are impressive numbers. Let them sink in. Emphasize that 87% is the AVERAGE across ALL algorithms -- individual results are even higher. The DFS path quality improvement is a nice unexpected finding to highlight.')
add_divider(doc)

# ============================================================
# SLIDE 13: AILS RESULTS
# ============================================================
add_slide_header(doc, 13, 'RESULTS: AILS EXPERIMENTS', 'DS2/DS3: Variable-size grids and topologies')

p = doc.add_paragraph()
run = p.add_run('Overall Performance (200x200, 25% density):')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'AILS-Base: 56.0% node reduction vs A* (Cohen\'s d = 0.82, large effect)')
add_bullet(doc, 'AILS-Adaptive: 51.4% node reduction (d = 0.76, medium-large effect)')
add_bullet(doc, 'Both highly significant: p < 0.001')

p = doc.add_paragraph()
run = p.add_run('Scalability (key finding):')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, '50x50: 5.1% node reduction (overhead dominates)')
add_bullet(doc, '200x200: 49.0% node reduction')
add_bullet(doc, '300x300: TIME-EFFICIENCY CROSSOVER -- AILS becomes FASTER than A*')
add_bullet(doc, '500x500: 76.8% node reduction')
add_bullet(doc, 'Node savings grow with grid size (corridor fraction shrinks quadratically)')

p = doc.add_paragraph()
run = p.add_run('Density Impact:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Sweet spot: 10-25% density -- AILS beats A* on both time and nodes')
add_bullet(doc, '10% density: AILS-Base faster (8.57ms vs 9.20ms) + 43.1% fewer nodes')
add_bullet(doc, '20% density: 29.6% faster + 57.4% fewer nodes')
add_bullet(doc, '>30% density: performance degrades, 40% density: success drops to 34%')

p = doc.add_paragraph()
run = p.add_run('Obstacle Patterns:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Random/Open: EXCELLENT (43-58% node reduction)')
add_bullet(doc, 'Clustered: POOR (70x slower -- large obstacle pockets)')
add_bullet(doc, 'Maze: POOR (81x slower -- winding corridors)')
add_bullet(doc, 'Room: WORST (116x slower -- narrow doorways)')

p = doc.add_paragraph()
run = p.add_run('[Visual: Scalability graph showing node reduction growing with grid size. Density chart. Topology comparison table.]')
run.italic = True
run.font.color.rgb = RGBColor(0, 0, 200)

add_note(doc, 'Key message: AILS works brilliantly on open/random layouts at moderate density. Be honest about where it fails -- the committee will appreciate the honest characterization.')
add_divider(doc)

# ============================================================
# SLIDE 14: ABLATION STUDY
# ============================================================
print("[8/12] Writing ablation study slides...", flush=True)
add_slide_header(doc, 14, 'RESULTS: ABLATION STUDY', 'Parameter sensitivity analysis')

p = doc.add_paragraph()
run = p.add_run('Radius Parameters (r_min, r_max):')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Wider corridors --> higher optimality but slower execution')
add_bullet(doc, '(1,5): 94.3% optimality, 8.2ms | (2,15): 100% optimality, 12.1ms')
add_bullet(doc, 'Default (2, ceil(0.1*min(H,W))): 99.8% optimality -- best balance')

p = doc.add_paragraph()
run = p.add_run('Window Half-Size (omega):')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Classic bias-variance tradeoff')
add_bullet(doc, 'Small window (3): noisy density estimates (45.2% improvement)')
add_bullet(doc, 'Sweet spot (7): optimal balance (62.2% improvement)')
add_bullet(doc, 'Large window (11): over-smoothed (59.1% improvement)')

p = doc.add_paragraph()
run = p.add_run('Density Sensitivity (alpha):')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'alpha=1.0: best optimality (98.7%), balanced corridor sizes')
add_bullet(doc, 'alpha<1: conservative (wide early), alpha>1: aggressive (narrow longer)')

p = doc.add_paragraph()
run = p.add_run('Strategy Comparison:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Base: 35.2% improvement, 89.4% optimality')
add_bullet(doc, 'Standard: 55.8% improvement, 96.7% optimality')
add_bullet(doc, 'Predictive: 62.2% improvement, 99.8% optimality <-- CLEAR WINNER')

p = doc.add_paragraph()
run = p.add_run('[Visual: Tables from ablation study. Highlight the sweet spots for each parameter.]')
run.italic = True
run.font.color.rgb = RGBColor(0, 0, 200)

add_note(doc, 'The ablation study shows that default parameters are well-chosen. Highlight the Predictive strategy as the standout result.')
add_divider(doc)

# ============================================================
# SLIDE 15: CROSS-STUDY DISCUSSION
# ============================================================
print("[9/12] Writing discussion slides...", flush=True)
add_slide_header(doc, 15, 'CROSS-STUDY DISCUSSION', 'ILS vs AILS: Complementary strengths')

p = doc.add_paragraph()
run = p.add_run('ILS Strengths:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Higher raw performance on uniform-density environments')
add_bullet(doc, '87.31% time reduction (larger than AILS\'s 51-56% node reduction)')
add_bullet(doc, 'Simpler implementation, lower overhead')
add_bullet(doc, 'Works with any algorithm (5 tested)')

p = doc.add_paragraph()
run = p.add_run('AILS Strengths:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Topological robustness across heterogeneous environments')
add_bullet(doc, 'Per-point adaptation prevents systematic over-expansion')
add_bullet(doc, 'Scales better with grid size (76.8% node reduction on 500x500)')
add_bullet(doc, 'Predictive strategy: 99.8% optimality')

p = doc.add_paragraph()
run = p.add_run('Complementary Nature:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'ILS: best choice for uniform-density, moderate-size grids')
add_bullet(doc, 'AILS: best choice for heterogeneous environments, large grids')
add_bullet(doc, 'Both: best at 10-25% density, random/open patterns')

p = doc.add_paragraph()
run = p.add_run('Answering the Research Questions:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'RQ1: ILS achieved 87.31% time reduction and 71.44% node reduction -- EXCEEDING RH1\'s 40-70% prediction')
add_bullet(doc, 'RQ2: AILS maintained efficiency on random/open patterns but degraded on structured topologies -- PARTIALLY supporting RH2')

add_note(doc, 'This is a crucial slide. Show that you understand how the two contributions relate and where each excels. The honest assessment of RH2 being "partially supported" shows scientific maturity.')
add_divider(doc)

# ============================================================
# SLIDE 16: COMPARISON WITH PRIOR METHODS
# ============================================================
add_slide_header(doc, 16, 'COMPARISON WITH PRIOR METHODS')

p = doc.add_paragraph()
run = p.add_run('vs. Jump Point Search (JPS):')
run.bold = True
run.font.size = Pt(12)
add_bullet(doc, 'JPS: 10-100x speedup but ONLY uniform-cost grids')
add_bullet(doc, 'ILS/AILS: more modest speedup but works with ANY cost model')
add_bullet(doc, 'Compatible: JPS could serve as base algorithm inside ILS corridor')

p = doc.add_paragraph()
run = p.add_run('vs. Hierarchical Methods (HPA*, Contraction Hierarchies):')
run.bold = True
run.font.size = Pt(12)
add_bullet(doc, 'Hierarchical: fast queries after expensive OFFLINE preprocessing')
add_bullet(doc, 'ILS/AILS: NO preprocessing needed -- suitable for dynamic environments')

p = doc.add_paragraph()
run = p.add_run('vs. Incremental Planners (D* Lite, LPA*):')
run.bold = True
run.font.size = Pt(12)
add_bullet(doc, 'Incremental: efficient repair across planning episodes')
add_bullet(doc, 'ILS/AILS: restrict search space within single query')
add_bullet(doc, 'COMPLEMENTARY: D* Lite inside AILS corridor = best of both worlds')

p = doc.add_paragraph()
run = p.add_run('vs. Theta*:')
run.bold = True
run.font.size = Pt(12)
add_bullet(doc, 'ILS borrows line-of-sight post-processing from Theta*')
add_bullet(doc, 'But ILS works as a wrapper on ANY algorithm, not just A*')

add_note(doc, 'Position your work clearly. You are not claiming to replace these methods -- you fill a specific gap and are complementary to existing techniques.')
add_divider(doc)

# ============================================================
# SLIDE 17: CONTRIBUTIONS & SIGNIFICANCE
# ============================================================
print("[10/12] Writing contributions and conclusion slides...", flush=True)
add_slide_header(doc, 17, 'CONTRIBUTIONS & SIGNIFICANCE')

p = doc.add_paragraph()
run = p.add_run('Contribution 1: Incremental Line Search (ILS) Framework')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 102, 0)

add_bullet(doc, 'First corridor-constrained search framework for non-uniform-cost grids')
add_bullet(doc, 'Algorithm-agnostic: works as wrapper around any classical search algorithm')
add_bullet(doc, '87.31% average execution time reduction, 71.44% node reduction')
add_bullet(doc, 'Preserves path optimality for optimal algorithms')
add_bullet(doc, 'Unexpected bonus: dramatically improves path quality for non-optimal algorithms')

p = doc.add_paragraph()
run = p.add_run('Contribution 2: Adaptive ILS (AILS) Framework')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 102, 0)

add_bullet(doc, 'Per-cell density-adaptive corridor using integral images')
add_bullet(doc, 'Three auto-selected strategies (Base, Standard, Predictive)')
add_bullet(doc, '76.8% node reduction on 500x500 grids')
add_bullet(doc, 'Predictive strategy: 62.2% time improvement, 99.8% optimality')
add_bullet(doc, 'Topological robustness across heterogeneous environments')

p = doc.add_paragraph()
run = p.add_run('Significance:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Enables real-time pathfinding on commodity hardware')
add_bullet(doc, 'No preprocessing, any cost model, any search algorithm')
add_bullet(doc, 'Applicable beyond biosecurity: warehouses, agriculture, SAR, general robotics')

add_note(doc, 'This is your showcase slide. Present the contributions with confidence. These are genuine, well-validated advances.')
add_divider(doc)

# ============================================================
# SLIDE 18: LIMITATIONS
# ============================================================
add_slide_header(doc, 18, 'LIMITATIONS')

add_bullet(doc, 'Overhead on small grids: ', bold_prefix='1. ')
doc.add_paragraph('   AILS slower than A* below ~300x300 due to corridor construction overhead')

add_bullet(doc, 'High-density environments: ', bold_prefix='2. ')
doc.add_paragraph('   Performance degrades above 30% density; poor on maze/room/clustered patterns')

add_bullet(doc, 'No formal sub-optimality bound: ', bold_prefix='3. ')
doc.add_paragraph('   99.8% empirical optimality but no worst-case (1+epsilon) guarantee')

add_bullet(doc, 'Synthetic benchmarks: ', bold_prefix='4. ')
doc.add_paragraph('   External validation on Moving AI benchmarks needed')

add_bullet(doc, 'Different hardware: ', bold_prefix='5. ')
doc.add_paragraph('   ILS (M1 Mac) and AILS (i7-12700K) experiments on different machines')

add_bullet(doc, 'Parameter dependence: ', bold_prefix='6. ')
doc.add_paragraph('   No automatic parameter tuning mechanism provided')

add_note(doc, 'Present limitations honestly. The committee respects self-awareness. Frame each limitation as an identified boundary, not a weakness.')
add_divider(doc)

# ============================================================
# SLIDE 19: FUTURE WORK
# ============================================================
add_slide_header(doc, 19, 'FUTURE WORK')

add_bullet(doc, 'Moving AI benchmark validation (external credibility)')
add_bullet(doc, 'C++ reimplementation (push crossover to smaller grids, larger-scale testing)')
add_bullet(doc, 'Formal sub-optimality analysis (instance-specific or probabilistic bounds)')
add_bullet(doc, 'Integration with D* Lite (corridor + incremental replanning)')
add_bullet(doc, 'Multi-agent pathfinding extension (per-agent corridors with conflict resolution)')
add_bullet(doc, 'Learned corridor axis (ML-predicted reference line instead of Bresenham)')
add_bullet(doc, 'Hardware deployment (ROS integration, drone/robot field trials)')
add_bullet(doc, 'Automatic parameter tuning (online adaptation of r_min, r_max, alpha, omega)')
add_bullet(doc, '3D extension (3D Bresenham + tube corridor + 3D integral volume)')

add_note(doc, 'Show that this work opens doors for future research. The committee likes to see that you have thought about where the field goes next.')
add_divider(doc)

# ============================================================
# SLIDE 20: CONCLUSION
# ============================================================
add_slide_header(doc, 20, 'CONCLUSION')

p = doc.add_paragraph()
run = p.add_run('This thesis developed two corridor-based pathfinding techniques for grid maps:')
run.bold = True

doc.add_paragraph()
add_bullet(doc, 'ILS: Corridor-constrained search achieving 87.31% time reduction and 71.44% node reduction across 5 algorithms, preserving optimality')
add_bullet(doc, 'AILS: Density-adaptive corridor achieving 76.8% node reduction on large grids with 99.8% optimality via Predictive strategy')
add_bullet(doc, 'The two methods are complementary: ILS for uniform-density, AILS for heterogeneous environments')
add_bullet(doc, 'Both work with any cost model, any search algorithm, with no preprocessing')
add_bullet(doc, 'Applicable to biosecurity, warehouse automation, agriculture, search-and-rescue, and general grid-based planning')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('In one sentence: ')
run.bold = True
p.add_run('By confining search to a narrow, optionally density-adaptive band around the straight line between start and goal, ILS and AILS achieve dramatic computational savings while maintaining path quality -- a simple idea with powerful results.')

add_note(doc, 'End with confidence. Summarize the key numbers one last time: 87% time reduction, 71% node reduction, 99.8% optimality. Thank the committee and invite questions.')
add_divider(doc)

# ============================================================
# SLIDE 21: THANK YOU / Q&A
# ============================================================
add_slide_header(doc, 21, 'THANK YOU & Q&A')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\nThank You\n\n')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Questions & Discussion')
run.font.size = Pt(16)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Elshahed Amr Moustafa Mohamed Aly Elsayed')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[Add your email/contact]')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

add_divider(doc)

# ============================================================
# APPENDIX: PRESENTATION TIPS
# ============================================================
print("[11/12] Writing presentation tips...", flush=True)
doc.add_page_break()
add_slide_header(doc, '', 'APPENDIX: PRESENTATION TIPS & TIMING GUIDE')

p = doc.add_paragraph()
run = p.add_run('Suggested Timing (45-minute presentation):')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Slides 1-2 (Title + Outline): 2 minutes')
add_bullet(doc, 'Slides 3-4 (Introduction + Biosecurity): 4 minutes')
add_bullet(doc, 'Slide 5 (Problem Statement): 3 minutes')
add_bullet(doc, 'Slide 6 (RQs + Objectives): 2 minutes')
add_bullet(doc, 'Slide 7 (Literature Review): 4 minutes')
add_bullet(doc, 'Slides 8-11 (Methodology): 10 minutes')
add_bullet(doc, 'Slides 12-14 (Results): 10 minutes')
add_bullet(doc, 'Slides 15-16 (Discussion + Comparison): 4 minutes')
add_bullet(doc, 'Slides 17-19 (Contributions + Limitations + Future): 4 minutes')
add_bullet(doc, 'Slide 20-21 (Conclusion + Q&A): 2 minutes')

p = doc.add_paragraph()
run = p.add_run('\nGeneral Tips:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'Use large, clear fonts on slides (minimum 24pt for body text)')
add_bullet(doc, 'One key message per slide -- do not overload')
add_bullet(doc, 'Use visuals wherever possible (diagrams, charts, animations)')
add_bullet(doc, 'Practice the presentation at least 3 times before the actual seminar')
add_bullet(doc, 'Prepare for 15-30 minutes of Q&A after the presentation')
add_bullet(doc, 'Have backup slides with detailed tables/data in case of specific questions')
add_bullet(doc, 'Be honest about limitations -- it shows scientific maturity')
add_bullet(doc, 'The Viva Q&A PDF document accompanies this file for detailed question preparation')

p = doc.add_paragraph()
run = p.add_run('\nKey Figures to Include from the Thesis:')
run.bold = True
run.font.size = Pt(12)

add_bullet(doc, 'figures/biosecurity_types_framework.png -- Biosecurity types and cybersecurity role')
add_bullet(doc, 'figures/framework_overview.png -- Three-pipeline research framework')
add_bullet(doc, 'figures/preprocessing.png -- Preprocessing pipeline flowchart')
add_bullet(doc, 'figures/strategy_flowchart.png -- AILS strategy selection flowchart')
add_bullet(doc, 'figures/ails_architecture.png -- AILS four-stage architecture')
add_bullet(doc, 'Create new: Corridor construction animation/diagram (ILS vs AILS)')
add_bullet(doc, 'Create new: Bar charts for ILS results by algorithm')
add_bullet(doc, 'Create new: Scalability line graph (node reduction vs grid size)')

# Save
print("[12/12] Saving document...", flush=True)
doc.save('/home/user/phd/Exit_Seminar_Slideshow_Content.docx')
print("Exit Seminar DOCX generated successfully!")
